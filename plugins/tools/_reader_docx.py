
import shutil
import subprocess


def _doc_to_docx(path: str):
    """Legacy .doc (OLE binary, readable by neither pandoc nor python-docx) → a LibreOffice-converted .docx temp copy.
    Returns the new path; None when soffice is unavailable or the conversion fails."""
    import os
    import tempfile
    if not shutil.which("soffice"):
        return None
    outdir = tempfile.mkdtemp(prefix="doc2docx_")
    env = dict(os.environ)
    env["SAL_USE_VCLPLUGIN"] = "svp"
    try:
        subprocess.run(["soffice", "--headless", "--convert-to", "docx", "--outdir", outdir, path],
                       capture_output=True, timeout=120, env=env)
    except Exception:
        return None
    new = os.path.join(outdir, os.path.splitext(os.path.basename(path))[0] + ".docx")
    return new if os.path.exists(new) else None


def _docx_to_md(path: str) -> str:
    """docx → markdown. The base path is pandoc → markdown
    (keeps run-level formatting signals: bold / italic / strikethrough / lists / tables, which is more than python-docx's plain text).
    When pandoc is unavailable it falls back to python-docx (text + heading levels only, run-level formatting lost).
    A legacy .doc is converted to .docx by LibreOffice first and then takes the same path.
    """
    if path.lower().endswith(".doc"):
        conv = _doc_to_docx(path)
        if not conv:
            return ("[read_file] .doc (legacy Word) detected but LibreOffice (soffice) is "
                    "unavailable for .doc→.docx conversion; cannot read.")
        path = conv
    if shutil.which("pandoc"):
        try:
            r = subprocess.run(
                ["pandoc", "-f", "docx", "-t", "markdown", "--wrap=none", path],
                capture_output=True, text=True, timeout=120,
            )
            if r.returncode == 0 and r.stdout.strip():
                return r.stdout
        except Exception:
            pass  # fall through to the python-docx fallback

    # Fallback: python-docx when pandoc is absent (plain text + headings)
    _ensure("docx", "python-docx")
    import docx
    d = docx.Document(path)
    out = []
    for p in d.paragraphs:
        if not p.text.strip():
            continue
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            lvl = "".join(c for c in style if c.isdigit()) or "1"
            out.append("#" * min(int(lvl), 6) + " " + p.text)
        else:
            out.append(p.text)
    for ti, t in enumerate(d.tables):
        out.append(f"\n**Table {ti + 1}:**")
        for ri, row in enumerate(t.rows):
            cells = [c.text for c in row.cells]
            out.append("| " + " | ".join(cells) + " |")
            if ri == 0:
                out.append("| " + " | ".join("---" for _ in cells) + " |")
    return "\n".join(out)
