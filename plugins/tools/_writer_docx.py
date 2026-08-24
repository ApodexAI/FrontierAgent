
# docx writing: incremental load-modify-save; the anchor is content text (find/after_text), not a positional index.
# Operations are implemented with python-docx; the feature set is informed by the docs_server of Mercor-Intelligence/archipelago (Apache-2.0),
# but the addressing and parameters are this project's own design (anchor text, resistant to index drift).
#
# Operations (op):
#   create(blocks=[...])                          build a new file from content blocks
#   replace_text(find, replace, count=0)          find/replace across the document (count=0 = all, otherwise the first N)
#   insert_paragraph(text, after_text=None, style=None, bold=False, italic=False)
#   insert_heading(text, level=1, after_text=None)
#   insert_table(rows=[[...]], after_text=None, header=True, style="Table Grid")
#   format_text(find, bold=None, italic=None, underline=None, font_size=None, font_color=None)
#   set_page_margins(top?, bottom?, left?, right?, section=0)        inches
#   set_page_orientation(orientation=portrait|landscape, section=0)
#   set_header_footer(header?, footer?, section=0)                   text; empty string = clear
#   add_image(image_path, after_text=None, width=None, height=None)  inches
# after_text=None → append at the end of the document; otherwise insert after "the first paragraph containing that text".
import os as _os


def _docx_mod():
    _ensure("docx", "python-docx")
    import docx
    return docx


_ALIGN = {"left": "LEFT", "center": "CENTER", "right": "RIGHT", "justify": "JUSTIFY"}


def _set_align(p, align):
    if not align:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    name = _ALIGN.get(str(align).lower())
    if name:
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, name)


def _add_hyperlink(p, url, text, *, bold=None, italic=None, color="0563C1", underline=True):
    """Append a real hyperlink run to paragraph p (python-docx has no direct API, so this goes through OXML)."""
    from docx.opc.constants import RELATIONSHIP_TYPE as _RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    r_id = p.part.relate_to(url, _RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), (color or "0563C1").lstrip("#"))
    rpr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    p._p.append(link)


def _apply_runs(p, text):
    """Write RichText (already normalised by _norm_runs) into paragraph p: formatting per run; a run with a link becomes a real hyperlink."""
    from docx.shared import Pt, RGBColor
    for rd in _norm_runs(text):
        if rd.get("link"):
            _add_hyperlink(p, rd["link"], rd.get("text", ""),
                           bold=rd.get("bold"), italic=rd.get("italic"),
                           color=(rd.get("color") or "0563C1"))
            continue
        r = p.add_run(rd.get("text", ""))
        if rd.get("bold") is not None:
            r.bold = bool(rd["bold"])
        if rd.get("italic") is not None:
            r.italic = bool(rd["italic"])
        if rd.get("underline") is not None:
            r.underline = bool(rd["underline"])
        if rd.get("strike") is not None:
            r.font.strike = bool(rd["strike"])
        if rd.get("color"):
            r.font.color.rgb = RGBColor.from_string(str(rd["color"]).lstrip("#"))
        if rd.get("size"):
            r.font.size = Pt(float(rd["size"]))
        if rd.get("font"):
            r.font.name = rd["font"]
            # python-docx only writes w:ascii / w:hAnsi — CJK characters use w:eastAsia, so
            # without setting it
            # "font":"SimSun" has no effect at all on Chinese/Japanese/Korean text (the
            # document default or theme font is still used, and Word shows
            # something other than intended). Write eastAsia too, so the requested font applies to every character.
            from docx.oxml.ns import qn as _qn
            r._element.rPr.rFonts.set(_qn("w:eastAsia"), rd["font"])


def _set_list(p, spec):
    """Paragraph list level (explicit k-v, never `- ` / `1.` inside text). spec={type:bullet|number, level:0-8}.
    level 0 uses the built-in List Bullet / Number styles; deeper levels fall back to the style plus an approximate left indent."""
    if not spec:
        return
    from docx.shared import Inches
    typ = str(spec.get("type", "bullet")).lower()
    lvl = int(spec.get("level", 0))
    base = "List Number" if typ.startswith("num") else "List Bullet"
    style = base if lvl == 0 else f"{base} {min(lvl + 1, 3)}"
    try:
        p.style = style
    except KeyError:
        p.style = base
    if lvl > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (lvl + 1))


def _set_para_flow(p, b):
    """Pagination properties: page_break_before / keep_with_next (paragraph level)."""
    pf = p.paragraph_format
    if b.get("page_break_before") is not None:
        pf.page_break_before = bool(b["page_break_before"])
    if b.get("keep_with_next") is not None:
        pf.keep_with_next = bool(b["keep_with_next"])


def _set_para_spacing(p, b):
    """Paragraph line spacing and space before/after: line_spacing (multiple) / space_before / space_after (points)."""
    from docx.shared import Pt
    pf = p.paragraph_format
    if b.get("line_spacing") is not None:
        pf.line_spacing = float(b["line_spacing"])
    if b.get("space_before") is not None:
        pf.space_before = Pt(float(b["space_before"]))
    if b.get("space_after") is not None:
        pf.space_after = Pt(float(b["space_after"]))


def _cell_fill(cell, color):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), str(color).lstrip("#"))
    tcPr.append(shd)


def _docx_cell(cell, val):
    """Write a table cell: val may be RichText (str | list of runs) or
    {content:RichText, bold?, align?, fill_color?}。"""
    content, extra = val, {}
    if isinstance(val, dict) and "content" in val:
        content, extra = val["content"], val
    p = cell.paragraphs[0]
    _apply_runs(p, content)
    if extra.get("bold"):
        for r in p.runs:
            r.bold = True
    if extra.get("align"):
        _set_align(p, extra["align"])
    if extra.get("fill_color"):
        _cell_fill(cell, extra["fill_color"])


def _header_row(tbl):
    """Bold the first row and mark it as a header row (`w:tblHeader`): repeats at the top of every page across a page break, and is semantically clear."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for c in tbl.rows[0].cells:
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _apply_col_widths(tbl, widths_in):
    """Set fixed column widths in inches. LibreOffice/Word default to auto-fit and ignore column widths, so all four are required:
    (1) table layout set to fixed (w:tblLayout type=fixed) (2) write w:tblGrid gridCol (twips,
    which is mainly what LibreOffice reads) (3) write tcW on every cell too (what Word reads) (4) turn autofit off."""
    if not widths_in:
        return
    from docx.oxml.ns import qn
    from docx.shared import Inches
    # (1)(4) The autofit=False setter inserts w:tblLayout (type=fixed) at the correct place in
    # tblPr per xsd:sequence (before tblLook). Do not hand-roll remove/append — append moves
    # tblLayout after tblLook, which is an invalid order and makes Word call the file corrupt (only LibreOffice is lenient).
    tbl.autofit = False
    ncols = len(tbl.columns)
    twips = []
    for w in widths_in[:ncols]:
        try:
            twips.append(max(round(float(w) * 1440), 1))
        except (TypeError, ValueError):
            twips.append(None)
    # ② tblGrid gridCol
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        gcols = grid.findall(qn("w:gridCol"))
        for j, tw in enumerate(twips):
            if tw is not None and j < len(gcols):
                gcols[j].set(qn("w:w"), str(tw))
    # (3) tcW per cell (iterate rows, avoiding the merged-cell trap in the columns iterator)
    for row in tbl.rows:
        for j, tw in enumerate(twips):
            if tw is not None and j < len(row.cells):
                row.cells[j].width = Inches(float(widths_in[j]))


def _iter_paras(doc):
    """Every paragraph in the body and inside table cells (find/replace has to cover tables)."""
    from docx.document import Document as _Doc
    yield from doc.paragraphs
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    _ = _Doc


def _replace_in_paragraph(p, find, replace, budget):
    """Run-level find/replace inside one paragraph (preserving formatting where possible). budget=[replacements left]; returns the number replaced."""
    if budget[0] == 0 or find not in p.text:
        return 0
    # The simple robust approach: rewrite the whole paragraph into its first run and clear the rest (paragraph-level formatting survives, run-level is lost — acceptable for v0)
    n = p.text.count(find)
    if budget[0] > 0:
        n = min(n, budget[0])
    new = p.text.replace(find, replace, n if budget[0] > 0 else -1)
    if not p.runs:
        p.add_run(new)
    else:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    if budget[0] > 0:
        budget[0] -= n
    return n


def _find_anchor_para(doc, after_text):
    for p in doc.paragraphs:
        if after_text in p.text:
            return p
    return None


def _insert_para_after(anchor_p, doc):
    """Insert a new empty paragraph after anchor_p and return it (via XML addnext). anchor=None → a new paragraph at the end."""
    if anchor_p is None:
        return doc.add_paragraph()
    new_p = anchor_p.insert_paragraph_before()  # a placeholder, moved after the anchor immediately below
    anchor_p._p.addnext(new_p._p)
    return new_p


def _add_block(doc, b):
    """Append one content block to the end of doc (used by create). Block schema is documented at the top of the file.
    Returns the type written (singular label, for the receipt count) or None (unknown block with no content → the caller fails loudly)."""
    t = b.get("type", "paragraph")
    if t == "heading":
        h = doc.add_heading("", level=int(b.get("level", 1)))
        _apply_runs(h, b.get("text", ""))
        _set_align(h, b.get("align"))
        return "heading"
    elif t == "paragraph":
        p = doc.add_paragraph(style=b.get("style"))
        _apply_runs(p, b.get("text", ""))
        _set_list(p, b.get("list"))
        _set_align(p, b.get("align"))
        _set_para_flow(p, b)
        _set_para_spacing(p, b)
        return "paragraph"
    elif t in ("bullet_list", "numbered_list", "list"):
        # "list" is an alias models often use; ordered/numbered=True or type=numbered_list → numbered, otherwise bulleted
        items = b.get("items", [])
        if not items:
            return None
        numbered = (t == "numbered_list") or bool(b.get("ordered") or b.get("numbered"))
        default_style = "List Number" if numbered else "List Bullet"
        for it in items:
            p = doc.add_paragraph(style=b.get("style", default_style))
            _apply_runs(p, it)  # RichText items supported
        return "list"
    elif t == "table":
        rows = b.get("rows", [])
        if not rows:
            return None
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        tbl.style = b.get("style", "Table Grid")
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                _docx_cell(tbl.cell(i, j), val)
        if b.get("header", True):
            _header_row(tbl)
        _apply_col_widths(tbl, b.get("column_widths_in"))
        return "table"
    elif t == "image":
        # The image block the create contract promises (dispatch used to be missing it: an image-only document reported "wrote nothing")
        import os as _os2

        from docx.shared import Inches
        ipath = b.get("path") or b.get("image_path")
        if not ipath or not _os2.path.exists(str(ipath)):
            return None  # path missing or nonexistent → counted as a bad block, create warns, nothing is silent
        kw = {}
        if b.get("width_in"):
            kw["width"] = Inches(float(b["width_in"]))
        if b.get("height_in"):
            kw["height"] = Inches(float(b["height_in"]))
        doc.add_picture(str(ipath), **kw)
        return "image"
    elif t == "page_break":
        doc.add_page_break()
        return "page_break"
    elif b.get("items"):
        # Unknown block but it carries items → treat as a list; never silently drop content
        for it in b.get("items", []):
            p = doc.add_paragraph(style="List Bullet")
            _apply_runs(p, it)
        return "list"
    elif b.get("text"):
        p = doc.add_paragraph()
        _apply_runs(p, b.get("text", ""))
        return "paragraph"
    return None  # unknown block with no text/items: do not silently write an empty paragraph, let create fail loudly


def _docx_write(path, op, args):
    docx = _docx_mod()
    from docx.shared import Pt, RGBColor

    if op == "create":
        if _os.path.exists(path) and not args.get("overwrite"):
            return _res(f"create refused: {path} already exists — use insert_*/replace_text to "
                        "add, or pass overwrite:true to rebuild", ok=False)
        doc = docx.Document()
        meta = args.get("metadata") or {}
        cp = doc.core_properties
        for k in ("title", "subject", "author", "comments"):
            if meta.get(k):
                setattr(cp, k, meta[k])
        blocks = args.get("blocks", [])
        counts, bad = {}, []
        for bi, b in enumerate(blocks):
            kind = _add_block(doc, b)
            if kind is None:
                bad.append(bi)
            else:
                counts[kind] = counts.get(kind, 0) + 1
        if not counts:
            return _res(f"create wrote nothing to {path}: {len(blocks)} block(s), all "
                        "empty/unknown — check block schema (type + text/items/rows)", ok=False)
        _os.makedirs(_os.path.dirname(path) or ".", exist_ok=True)
        doc.save(path)
        warn = f"block(s) {bad} wrote nothing (check schema)" if bad else None
        return _res(f"created docx: {path}", counts=counts, warn=warn)

    if not _os.path.exists(path):
        return f"[error] file not found (edit needs existing file): {path}"
    doc = docx.Document(path)

    if op == "replace_text":
        find, repl = args["find"], args.get("replace", "")
        budget = [int(args.get("count", 0)) or -1]
        budget[0] = budget[0] if budget[0] > 0 else -1
        cnt = 0
        for p in _iter_paras(doc):
            if budget[0] == 0:
                break
            cnt += _replace_in_paragraph(p, find, repl, budget)
        doc.save(path)
        return _res(f"replaced {cnt} occurrence(s) of {find!r}",
                    warn=(f"0 matches for {find!r}" if cnt == 0 else None))

    if op == "insert_paragraph":
        if args.get("after_text") and _find_anchor_para(doc, args["after_text"]) is None:
            return _res("insert_paragraph skipped",
                        warn=f"anchor {args['after_text']!r} not found — 0 inserted")
        anchor = _find_anchor_para(doc, args["after_text"]) if args.get("after_text") else None
        p = _insert_para_after(anchor, doc)
        if args.get("style"):
            p.style = args["style"]
        _apply_runs(p, args.get("text", ""))
        _set_list(p, args.get("list"))
        _set_align(p, args.get("align"))
        _set_para_flow(p, args)
        _set_para_spacing(p, args)
        doc.save(path)
        return f"inserted paragraph ({'after anchor' if anchor else 'at end'})"

    if op == "insert_heading":
        lvl = int(args.get("level", 1))
        if args.get("after_text"):
            anchor = _find_anchor_para(doc, args["after_text"])
            if anchor is None:
                return _res("insert_heading skipped",
                            warn=f"anchor {args['after_text']!r} not found — 0 inserted")
            p = _insert_para_after(anchor, doc)
            p.style = doc.styles[f"Heading {min(lvl, 9)}"]
            _apply_runs(p, args["text"])
        else:
            h = doc.add_heading("", level=lvl)
            _apply_runs(h, args["text"])
        doc.save(path)
        return f"inserted heading L{lvl}"

    if op == "insert_table":
        rows = args["rows"]
        anchor = _find_anchor_para(doc, args["after_text"]) if args.get("after_text") else None
        if args.get("after_text") and anchor is None:
            return _res("insert_table skipped",
                        warn=f"anchor {args['after_text']!r} not found — 0 inserted")
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        tbl.style = args.get("style", "Table Grid")
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                _docx_cell(tbl.cell(i, j), val)
        if args.get("header", True):
            _header_row(tbl)
        _apply_col_widths(tbl, args.get("column_widths_in"))
        if args.get("cant_split"):      # stop a table row from splitting across pages (w:cantSplit)
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            for row in tbl.rows:
                trPr = row._tr.get_or_add_trPr()
                # Insert before trHeight/tblHeader per CT_TrPr's xsd:sequence (a header row has
                # already appended w:tblHeader, and appending again would invert the order and be invalid); also guards against double insertion.
                if trPr.find(qn("w:cantSplit")) is None:
                    trPr.insert_element_before(
                        OxmlElement("w:cantSplit"),
                        "w:trHeight", "w:tblHeader", "w:tblCellSpacing", "w:jc", "w:hidden")
        if anchor is not None:
            anchor._p.addnext(tbl._tbl)
        doc.save(path)
        return f"inserted table {len(rows)}×{len(rows[0])}"

    if op == "format_text":
        find = args["find"]
        hits = 0
        for p in _iter_paras(doc):
            if find not in p.text:
                continue
            # Run level: apply formatting to runs that fall entirely inside the matched text (simplification: a whole-paragraph match formats the whole paragraph)
            for r in p.runs:
                if (r.text and r.text in find) or find in (r.text or ""):
                    if args.get("bold") is not None:
                        r.bold = bool(args["bold"])
                    if args.get("italic") is not None:
                        r.italic = bool(args["italic"])
                    if args.get("underline") is not None:
                        r.underline = bool(args["underline"])
                    if args.get("strike") is not None:
                        r.font.strike = bool(args["strike"])
                    sz = args.get("font_size") or args.get("size")  # size/color = the parameter names the documented contract uses
                    if sz:
                        r.font.size = Pt(float(sz))
                    col = args.get("font_color") or args.get("color")
                    if col:
                        r.font.color.rgb = RGBColor.from_string(str(col).lstrip("#"))
                    hits += 1
        doc.save(path)
        return _res(f"formatted {hits} run(s) matching {find!r}",
                    warn=(f"0 runs matched {find!r}" if hits == 0 else None))

    if op == "set_page_margins":
        from docx.shared import Inches
        sec = doc.sections[int(args.get("section", 0))]
        for side in ("top", "bottom", "left", "right"):
            if args.get(side) is not None:
                setattr(sec, f"{side}_margin", Inches(float(args[side])))
        doc.save(path)
        return f"set page margins on section {args.get('section', 0)}"

    if op == "set_page_orientation":
        from docx.enum.section import WD_ORIENT
        sec = doc.sections[int(args.get("section", 0))]
        want = str(args.get("orientation", "portrait")).lower()
        is_land = want.startswith("land")
        sec.orientation = WD_ORIENT.LANDSCAPE if is_land else WD_ORIENT.PORTRAIT
        # python-docx does not swap width/height automatically, so swap them by hand for this to take effect
        w, h = sec.page_width, sec.page_height
        if (is_land and w < h) or ((not is_land) and w > h):
            sec.page_width, sec.page_height = h, w
        doc.save(path)
        return f"set orientation={want} on section {args.get('section', 0)}"

    if op == "set_header_footer":
        sec = doc.sections[int(args.get("section", 0))]
        if args.get("header") is not None:
            sec.header.is_linked_to_previous = False
            hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
            hp.text = args["header"]
        if args.get("footer") is not None:
            sec.footer.is_linked_to_previous = False
            fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
            fp.text = args["footer"]
        doc.save(path)
        return f"set header/footer on section {args.get('section', 0)}"

    if op == "add_image":
        from docx.shared import Inches
        kw = {}
        if args.get("width"):
            kw["width"] = Inches(float(args["width"]))
        if args.get("height"):
            kw["height"] = Inches(float(args["height"]))
        if args.get("after_text"):
            anchor = _find_anchor_para(doc, args["after_text"])
            if anchor is None:
                return _res("add_image skipped",
                            warn=f"anchor {args['after_text']!r} not found — image not added")
            p = _insert_para_after(anchor, doc)
            p.add_run().add_picture(args["image_path"], **kw)
        else:
            doc.add_picture(args["image_path"], **kw)
        doc.save(path)
        return f"added image {_os.path.basename(args['image_path'])}"

    if op == "add_hyperlink":
        find, url = args["find"], args["url"]
        for p in _iter_paras(doc):
            if find in p.text:
                full = p.text
                i = full.find(find)
                before, after = full[:i], full[i + len(find):]
                for r in list(p.runs):
                    r.text = ""
                if before:
                    p.add_run(before)
                _add_hyperlink(p, url, find)
                if after:
                    p.add_run(after)
                doc.save(path)
                return f"added hyperlink on {find!r} → {url}"
        return _res("add_hyperlink skipped", warn=f"text not found for hyperlink: {find!r}")

    if op == "set_page_number":
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        loc = str(args.get("location", "footer")).lower()
        sec = doc.sections[0]
        container = sec.header if loc == "header" else sec.footer
        container.is_linked_to_previous = False
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        for r in list(p.runs):
            r.text = ""
        _set_align(p, args.get("align", "center"))

        def _field(instr):
            run = p.add_run()
            fc1 = OxmlElement("w:fldChar")
            fc1.set(qn("w:fldCharType"), "begin")
            run._r.append(fc1)
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = instr
            run._r.append(it)
            fc2 = OxmlElement("w:fldChar")
            fc2.set(qn("w:fldCharType"), "end")
            run._r.append(fc2)

        if args.get("of_total"):
            _field("PAGE")
            p.add_run(" / ")
            _field("NUMPAGES")
        else:
            _field("PAGE")
        if args.get("start") is not None or args.get("fmt"):
            sectPr = sec._sectPr
            pg = sectPr.find(qn("w:pgNumType"))
            if pg is None:
                pg = OxmlElement("w:pgNumType")
                sectPr.append(pg)
            if args.get("start") is not None:
                pg.set(qn("w:start"), str(int(args["start"])))
            fmtmap = {"decimal": "decimal", "roman_lower": "lowerRoman",
                      "roman_upper": "upperRoman", "alpha": "lowerLetter"}
            if args.get("fmt"):
                pg.set(qn("w:fmt"), fmtmap.get(args["fmt"], "decimal"))
        doc.save(path)
        return f"set page number in {loc} (of_total={bool(args.get('of_total'))})"

    if op == "format_paragraph":
        # Adjust an existing paragraph's spacing / alignment / pagination (anchor = the find text).
        find = args.get("find")
        if not find:
            return "[error] format_paragraph needs 'find' (anchor text)"
        n = 0
        for p in _iter_paras(doc):
            if find not in p.text:
                continue
            _set_para_spacing(p, args)
            _set_align(p, args.get("align"))
            _set_para_flow(p, args)
            n += 1
        doc.save(path)
        return _res(f"formatted {n} paragraph(s) matching {find!r}",
                    warn=(f"0 matched {find!r}" if n == 0 else None))

    return f"[error] unknown docx op: {op}"
